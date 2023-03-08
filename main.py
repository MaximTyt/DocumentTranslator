import shutil
from googletrans import Translator
from nltk import tokenize, download
from docx.api import Document
import os
from win32com import client as wc
import pythoncom
from pdf2docx import Converter
import openpyxl
len_max = 3000
# download('punkt')
LANGUAGES = {
    'af': 'Африканский',
    'sq': 'Албанский',
    'am': 'Амхарский',
    'ar': 'Арабский',
    'hy': 'Армянский',
    'az': 'Азербайджанский',
    'eu': 'Баскский',
    'be': 'Белорусский',
    'bn': 'Бенгальский',
    'bs': 'Боснийский',
    'bg': 'Болгарский',
    'ca': 'Каталанский',
    'ceb': 'Себуанский',
    'ny': 'Чева',
    'zh-cn': 'Китайский (упрощённый)',
    'zh-tw': 'Китайский (традиционный)',
    'co': 'Корсиканский',
    'hr': 'Хорватский',
    'cs': 'Чешский',
    'da': 'Датский',
    'nl': 'Нидерландский',
    'en': 'Английский',
    'eo': 'Эсперанто',
    'et': 'Эстонсикй',
    'tl': 'Филлипинский',
    'fi': 'Финский',
    'fr': 'Французский',
    'fy': 'Фризский',
    'gl': 'Галисийский',
    'ka': 'Грузинский',
    'de': 'Немецкий',
    'el': 'Греческий',
    'gu': 'Гуджарати',
    'ht': 'Креольский (гаити)',
    'ha': 'Хауса',
    'haw': 'Гавайский',
    'he': 'Иврит',
    'hi': 'Хинди',
    'hmn': 'Хмонг',
    'hu': 'Венгерский',
    'is': 'Исландский',
    'ig': 'Игбо',
    'id': 'Индонезисйский',
    'ga': 'Ирландский',
    'it': 'Итальянский',
    'ja': 'Японский',
    'jw': 'Яванский',
    'kn': 'Каннада',
    'kk': 'Казахский',
    'km': 'Кхмерский',
    'ko': 'Корейский',
    'ku': 'Курдский (курманджи)',
    'ky': 'Киргизский',
    'lo': 'Лаосский',
    'la': 'Латинский',
    'lv': 'Латышский',
    'lt': 'Литовский',
    'lb': 'Люксембургский',
    'mk': 'Македонский',
    'mg': 'Малагасийский',
    'ms': 'Малайский',
    'ml': 'Малаялам',
    'mt': 'Мальтийский',
    'mi': 'Маори',
    'mr': 'Маратхи',
    'mn': 'Монгольский',
    'my': 'Мейтейлон (Манипури)',
    'ne': 'Непальский',
    'no': 'Норвежский',
    'or': 'Ория',
    'ps': 'Пушту',
    'fa': 'Персидский',
    'pl': 'Польский',
    'pt': 'Португальский',
    'pa': 'Панджаби',
    'ro': 'Румынский',
    'ru': 'Русский',
    'sm': 'Самоанский',
    'gd': 'Шотландский (Гэльский)',
    'sr': 'Сербский',
    'st': 'Сесото',
    'sn': 'Шона',
    'sd': 'Синдхи',
    'si': 'Сингальский',
    'sk': 'Словацкий',
    'sl': 'словенский',
    'so': 'Сомалийский',
    'es': 'Испанский',
    'su': 'Сунданский',
    'sw': 'Суахили',
    'sv': 'Шведский',
    'tg': 'Таджикский',
    'ta': 'Тамильский',
    'te': 'Телугу',
    'th': 'Тайский',
    'tr': 'Турецкий',
    'uk': 'Украинский',
    'ur': 'Урду',
    'ug': 'Уйгурский',
    'uz': 'Узбекский',
    'vi': 'Вьетнамский',
    'cy': 'Валлийский',
    'xh': 'Коса',
    'yi': 'Идиш',
    'yo': 'Йоруба',
    'zu': 'Зулу'}

# FileFormat = 11 .doc, 12 .docx, 17 .pdf, 51 .xlsx, 56 .xls
def convertXlsAndXlsx(file_path, src_ext, dest_ext, file_format):
    file_path = os.path.abspath(file_path)
    joinedPath = file_path.replace(src_ext, dest_ext)
    pythoncom.CoInitialize()
    excel = wc.gencache.EnsureDispatch('Excel.Application')
    wb = excel.Workbooks.Open(file_path)
    wb.SaveAs(joinedPath, FileFormat=file_format)
    wb.Close()
    excel.Application.Quit()
    print(f'Конвертация {file_path} -> {joinedPath} прошла успешно!')
    return joinedPath


def convertDocAndDocx(file_path, src_ext, dest_ext, file_format, type_app):
    file_path = os.path.abspath(file_path)
    joinedPath = file_path.replace(src_ext, dest_ext)
    pythoncom.CoInitialize()
    w = wc.Dispatch(type_app)
    if file_format in (51, 56):
        doc = w.Workbooks.Open(file_path)
    else:
        doc = w.Documents.Open(file_path)
    doc.SaveAs(joinedPath, FileFormat=file_format)
    doc.Close()
    w.Quit()
    print(f'Конвертация {file_path} -> {joinedPath} прошла успешно!')
    return joinedPath

def convertPDFtoDocx(path):
    try:
        cv_obj = Converter(path)
        path = path.replace('.pdf', '.docx')
        cv_obj.convert(path)
        cv_obj.close()
    except:
        print('Conversion Failed')
    else:
        return path


def translateText(all_text, dest_lang, src_lang):
    new_all_text = []
    for i in range(len(all_text)):
        if '\n' in all_text[i] and len(all_text[i]) > 1:
            temp = all_text[i].split('\n')
            new_all_text.extend(temp)
        elif len(all_text[i]) > len_max:
            temp = tokenize.sent_tokenize(all_text[i])
            new_all_text.extend(temp)
        elif all_text[i] != '':
            new_all_text.append(all_text[i])
        if '' in new_all_text:
            new_all_text.remove('')
    all_text = sorted(list(set(new_all_text)), reverse=True, key=len)
    if '' in all_text:
        all_text.remove('')
    block_list = []
    temp = ''
    for i in all_text:
        if len(temp + i + '\n') <= len_max:
            temp += i + '\n'
        else:
            block_list.append(temp[:-1])
            temp = i + '\n'
    if temp != '':
        block_list.append(temp[:-1])
    new_all_text = []
    translator = Translator()
    result = [translator.translate(i, dest=dest_lang, src=src_lang).text for i in block_list]
    for i in result:
        temp = i.split('\n')
        new_all_text.extend(temp)
    return all_text, new_all_text


def translate_txt(path, dest_lang, src_lang):
    text = ''
    with open(path, encoding="utf_8") as f:
        text = f.read()
    all_text = text.split('\n')
    all_text, new_all_text = translateText(all_text, dest_lang, src_lang)
    text_len = min(len(all_text), len(new_all_text))
    for i in range(text_len):
        text = text.replace(all_text[i], new_all_text[i])
    path = path.replace('upload_files', 'translated_upload_files')
    with open(path, 'w', encoding="utf-8") as f:
        f.write(text)
    f.close()
    shutil.rmtree('upload_files')


def translate_doc_docx_pdf(path, dest_lang, src_lang):
    origin_path = path
    if path.endswith('.doc'):
        path = convertDocAndDocx(path, '.doc', '.docx', 12, 'Word.Application')
    elif path.endswith('pdf'):
        path = convertPDFtoDocx(path)
    document = Document(path)
    all_text = []
    for p in document.paragraphs:
        all_text.append(p.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)

    all_text, new_all_text = translateText(all_text, dest_lang, src_lang)

    text_len = min(len(all_text), len(new_all_text))
    for i in range(text_len):
        for paragraph in document.paragraphs:
            if all_text[i] in paragraph.text:
                paragraph.text = paragraph.text.replace(all_text[i], new_all_text[i])
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if all_text[i] in paragraph.text:
                            paragraph.text = paragraph.text.replace(all_text[i], new_all_text[i])
    path = path.replace('upload_files', 'translated_upload_files').replace('.DOCX', '.docx')
    document.save(path)
    if origin_path.lower().endswith('.doc'):
        convertDocAndDocx(path, '.docx', '.doc', 11, 'Word.Application')
    if origin_path.lower().endswith('.pdf'):
        convertDocAndDocx(path, '.docx', '.pdf', 17, 'Word.Application')

def translate_xls_xlsx(path, dest_lang, src_lang):
    origin_path = path
    all_text = []
    if path.endswith('.xls'):
        path = convertXlsAndXlsx(path, '.xls', '.xlsx', 51)
    wb = openpyxl.load_workbook(path)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is not None:
                    all_text.append(cell.value)
    all_text, new_all_text = translateText(all_text, dest_lang, src_lang)
    text_len = min(len(all_text), len(new_all_text))
    for i in range(text_len):
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        if all_text[i] in cell.value:
                            cell.value = cell.value.replace(all_text[i], new_all_text[i])
    path = path.replace('upload_files', 'translated_upload_files')
    wb.save(path)
    if origin_path.endswith('.xls'):
        convertXlsAndXlsx(path, '.xlsx', '.xls', 56)
